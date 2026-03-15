"""
Y2K Global AI Portfolio - DOCX Generator v4
Marketing-agency quality, dark theme, modern typography
"""
import os
import copy
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from lxml import etree

BASE = r"C:\git\y2k-portfolio"
OUT = os.path.join(BASE, "Y2K-AI-Portfolio.docx")
LOGO_W = os.path.join(BASE, "y2klogowhite.png")
BANNER = os.path.join(BASE, "banner.png")

# Colors
DARK = "0B1120"
PANEL = "131D33"
ACCENT = "00C8F0"
ACCENT_DIM = "0090A8"
TEXT_COL = "E8E8E8"
DIM_COL = "AAAAAA"
MUTED_COL = "777777"
WHITE = "FFFFFF"

doc = Document()

# ── Global styles ──
style = doc.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)
style.font.color.rgb = RGBColor.from_string(TEXT_COL)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.line_spacing = 1.15

# ── Page setup: Landscape A4, narrow margins ──
section = doc.sections[0]
section.orientation = WD_ORIENT.LANDSCAPE
section.page_width = Cm(29.7)
section.page_height = Cm(21.0)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(1.8)
section.bottom_margin = Cm(1.5)

# ── Dark background ──
bg_elem = parse_xml(f'<w:background {nsdecls("w")} w:color="{DARK}"/>')
doc.element.insert(0, bg_elem)
settings = doc.settings.element
display_bg = parse_xml(f'<w:displayBackgroundShape {nsdecls("w")}/>')
settings.append(display_bg)

# ── Helpers ──

def rgb(hex_str):
    return RGBColor.from_string(hex_str)

def add_spacer(height_pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    run.font.size = Pt(height_pt)
    return p

def add_text(text, size=11, color=TEXT_COL, bold=False, italic=False, align=None, space_after=0, space_before=0):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Calibri'
    return p

def add_multirun(runs_list, align=None, space_after=0, space_before=0, line_spacing=1.2):
    """runs_list: [(text, size, color, bold, italic), ...]"""
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = line_spacing
    for text, size, color, bold, italic in runs_list:
        run = p.add_run(text)
        run.font.size = Pt(size)
        run.font.color.rgb = rgb(color)
        run.font.bold = bold
        run.font.italic = italic
        run.font.name = 'Calibri'
    return p

def add_accent_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    # Use a thin horizontal rule via bottom border on paragraph
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="6" w:space="1" w:color="{ACCENT}"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    return p

def add_muted_line():
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="1" w:color="333344"/>'
        f'</w:pBdr>'
    )
    pPr.append(borders)
    return p

def add_section_title(text, size=24):
    p = add_text(text, size=size, color=ACCENT, bold=True, space_before=4)
    add_accent_line()
    return p

def add_page_break():
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.enum.text.WD_BREAK.PAGE)

def shade_cell(cell, color):
    """Apply background shading to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._element.get_or_add_tcPr().append(shading)

def styled_cell_text(cell, text, size=10, color=TEXT_COL, bold=False, italic=False, align=None):
    """Add styled text to a cell."""
    p = cell.paragraphs[0]
    p.clear()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.font.bold = bold
    run.font.italic = italic
    run.font.name = 'Calibri'
    return p

def add_card_table(rows_data, col_widths=None):
    """Create a dark panel table. rows_data: list of lists of (text, size, color, bold)"""
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Remove borders, add dark background
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    for i, row in enumerate(table.rows):
        for j, cell in enumerate(row.cells):
            shade_cell(cell, PANEL)
            text, size, color, bold = rows_data[i][j]
            styled_cell_text(cell, text, size, color, bold)
            # Cell margins
            tc = cell._element
            tcPr = tc.get_or_add_tcPr()
            margins = parse_xml(
                f'<w:tcMar {nsdecls("w")}>'
                f'  <w:top w:w="80" w:type="dxa"/>'
                f'  <w:left w:w="120" w:type="dxa"/>'
                f'  <w:bottom w:w="80" w:type="dxa"/>'
                f'  <w:right w:w="120" w:type="dxa"/>'
                f'</w:tcMar>'
            )
            tcPr.append(margins)
    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    return table

def tech_pill(text):
    """Return a formatted tech tag string."""
    return text

import docx.enum.text

# ════════════════════════════════════════
# PAGE 1: COVER
# ════════════════════════════════════════

# Try to add banner as background image in header
try:
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hp.add_run()
    run.add_picture(BANNER, width=Cm(29.7))
    # Make it behave as background (behind doc)
    inline = run._element.findall(qn('w:drawing'))[0]
    inline_elem = inline.find(qn('wp:inline'))
    if inline_elem is not None:
        # Convert inline to anchor with behindDoc
        anchor = copy.deepcopy(inline_elem)
        anchor.tag = qn('wp:anchor')
        anchor.set('behindDoc', '1')
        anchor.set('locked', '0')
        anchor.set('layoutInCell', '1')
        anchor.set('allowOverlap', '1')
        anchor.set('simplePos', '0')
        anchor.set('relativeHeight', '0')
        # Add position elements
        posH = etree.SubElement(anchor, qn('wp:positionH'))
        posH.set('relativeFrom', 'page')
        posHoff = etree.SubElement(posH, qn('wp:align'))
        posHoff.text = 'center'
        posV = etree.SubElement(anchor, qn('wp:positionV'))
        posV.set('relativeFrom', 'page')
        posVoff = etree.SubElement(posV, qn('wp:align'))
        posVoff.text = 'top'
        # Remove extent and replace
        inline.remove(inline_elem)
        inline.append(anchor)
except Exception as e:
    print(f"Banner note: {e}")

add_spacer(80)
add_spacer(40)

# Logo
try:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run()
    run.add_picture(LOGO_W, width=Cm(5))
except:
    pass

add_spacer(6)

# Title
add_text("Y2K Global", size=38, color=WHITE, bold=True, align='center', space_after=2)
add_accent_line()
add_spacer(4)
add_text("AI Portfolio", size=20, color=ACCENT, align='center', space_after=4)
add_text("AI Solutions & Consulting", size=13, color=DIM_COL, align='center', space_after=20)

add_spacer(40)

# Bottom info
add_text("Gergo Kiss, MSc  ·  Ing. Ramazan Yildirim", size=12, color=TEXT_COL, align='center', space_after=4)
add_text("+43 1 442 20 143  ·  office@y2k.global  ·  www.y2k.global", size=10, color=MUTED_COL, align='center')

# ════════════════════════════════════════
# PAGE 2: ABOUT + HOW WE WORK
# ════════════════════════════════════════
doc.add_page_break()

add_spacer(4)
add_section_title("About Y2K Global", 26)
add_spacer(6)

add_text(
    "Y2K Global is a technology consultancy that turns complex business challenges into intelligent, "
    "automated systems. We specialize in artificial intelligence, machine learning, and software "
    "optimization — delivering solutions that create measurable value, not just impressive demos.",
    size=11.5, color=TEXT_COL, space_after=8
)
add_text(
    "With deep expertise spanning natural language processing, predictive analytics, document intelligence, "
    "and full-stack development, we bridge the gap between cutting-edge AI research and production-ready "
    "business solutions.",
    size=11.5, color=TEXT_COL, space_after=14
)

# Leadership card
add_text("Leadership", size=16, color=ACCENT, bold=True, space_after=4)
add_muted_line()
add_spacer(4)

table = doc.add_table(rows=1, cols=1)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cell = table.rows[0].cells[0]
shade_cell(cell, PANEL)
p = cell.paragraphs[0]
run = p.add_run("Gergo Kiss, MSc  ·  Ing. Ramazan Yildirim")
run.font.size = Pt(13)
run.font.color.rgb = rgb(WHITE)
run.font.bold = True
run.font.name = 'Calibri'
p2 = cell.add_paragraph()
run2 = p2.add_run("Co-Founders")
run2.font.size = Pt(10)
run2.font.color.rgb = rgb(ACCENT)
run2.font.name = 'Calibri'
p3 = cell.add_paragraph()
p3.paragraph_format.space_before = Pt(4)
run3 = p3.add_run(
    "Gergo and Ramazan co-founded and jointly lead Y2K Global. As former fellow students during "
    "their Computer Science studies in Vienna, they developed a shared passion for applying artificial "
    "intelligence and machine learning to real-world business challenges."
)
run3.font.size = Pt(10)
run3.font.color.rgb = rgb(DIM_COL)
run3.font.name = 'Calibri'
# Cell padding
tc = cell._element
tcPr = tc.get_or_add_tcPr()
margins = parse_xml(
    f'<w:tcMar {nsdecls("w")}>'
    f'  <w:top w:w="140" w:type="dxa"/>'
    f'  <w:left w:w="200" w:type="dxa"/>'
    f'  <w:bottom w:w="140" w:type="dxa"/>'
    f'  <w:right w:w="200" w:type="dxa"/>'
    f'</w:tcMar>'
)
tcPr.append(margins)
# Remove table borders
tbl = table._tbl
tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
borders = parse_xml(
    f'<w:tblBorders {nsdecls("w")}>'
    f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
    f'</w:tblBorders>'
)
tblPr.append(borders)

add_spacer(12)

# How We Work
add_text("How We Work", size=16, color=ACCENT, bold=True, space_after=4)
add_muted_line()
add_spacer(6)

phases = [
    ("01", "Discover", "Deep immersion in your business context, data, and constraints. We identify where AI creates real value."),
    ("02", "Architect", "Systems that integrate with existing infrastructure, meet security requirements, and scale with your growth."),
    ("03", "Build & Validate", "Iterative development against real data. Every sprint produces working software you can evaluate."),
    ("04", "Deploy & Evolve", "Production deployment with monitoring and ongoing optimization as your needs change."),
]

for num, title, desc in phases:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Number cell
    cell0 = table.rows[0].cells[0]
    cell0.width = Cm(1.5)
    shade_cell(cell0, ACCENT)
    p = cell0.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(num)
    run.font.size = Pt(14)
    run.font.color.rgb = rgb(DARK)
    run.font.bold = True
    run.font.name = 'Calibri'
    # Content cell
    cell1 = table.rows[0].cells[1]
    shade_cell(cell1, PANEL)
    p = cell1.paragraphs[0]
    run = p.add_run(title)
    run.font.size = Pt(12)
    run.font.color.rgb = rgb(WHITE)
    run.font.bold = True
    run.font.name = 'Calibri'
    p2 = cell1.add_paragraph()
    run2 = p2.add_run(desc)
    run2.font.size = Pt(10)
    run2.font.color.rgb = rgb(DIM_COL)
    run2.font.name = 'Calibri'
    # Padding
    for cell in [cell0, cell1]:
        tc = cell._element
        tcPr = tc.get_or_add_tcPr()
        margins = parse_xml(
            f'<w:tcMar {nsdecls("w")}>'
            f'  <w:top w:w="100" w:type="dxa"/>'
            f'  <w:left w:w="140" w:type="dxa"/>'
            f'  <w:bottom w:w="100" w:type="dxa"/>'
            f'  <w:right w:w="140" w:type="dxa"/>'
            f'</w:tcMar>'
        )
        tcPr.append(margins)
    # Remove borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    add_spacer(3)

# ════════════════════════════════════════
# CASE STUDIES (pages 3-9)
# ════════════════════════════════════════

cases = [
    ("Funding & Grants Automation",
     "AI-Powered Funding Discovery & Application Platform",
     "European organizations spend hundreds of hours quarterly searching for relevant funding opportunities across fragmented EU and national databases. Applications are prepared manually, deadlines are missed, and success rates suffer from poor opportunity-fit matching.",
     "We developed an end-to-end AI platform that continuously monitors EU portals, national databases, and email communications for funding opportunities. The system uses NLP to evaluate each opportunity against the organization's profile, scores relevance and feasibility, generates comprehensive application briefs, and manages the full lifecycle from discovery through submission.",
     ["Automated scanning of 100+ funding sources", "AI-driven opportunity matching and scoring", "Intelligent application brief generation", "Full lifecycle tracking and management", "Significant reduction in manual research hours"],
     ["Python", "FastAPI", "PostgreSQL", "NLP", "Claude AI", "Email/IMAP", "Web Scraping"]),
    ("Private AI Infrastructure",
     "Private AI Platform with Standard SDK Integration",
     "Organizations needed to deploy AI and large language model capabilities within their own infrastructure — without sending sensitive data to external cloud providers. Existing solutions required deep AI expertise and custom integration for each use case.",
     "We built a private AI platform that runs entirely within the client's infrastructure, providing an OpenAI-compatible API as a standard SDK interface. Existing tools connect with zero code changes. The platform features intelligent multi-model routing, a comprehensive admin dashboard, and enterprise-grade access controls.",
     ["Fully private deployment — no data leaves the organization", "OpenAI-compatible SDK for seamless integration", "Multi-model routing for cost/performance optimization", "Admin dashboard for monitoring and governance", "Zero-friction adoption across departments"],
     ["Python", "LLM Inference", "API Gateway", "Docker", "REST API", "Admin UI"]),
    ("Enterprise System Optimization",
     "Fleet Information System — Performance & AI Enhancement",
     "A fleet management platform serving hundreds of locations suffered from critical performance bottlenecks. Key reports took over 14 seconds to load. Data type mismatches caused intermittent failures that eroded user trust.",
     "We performed a deep technical audit of the .NET/SQL Server stack, identified root causes in query design, indexing, and ORM configuration. Applied surgical optimizations: replaced expensive SQL functions, added covering indexes, implemented NoTracking for reads. Extended the system with an AI companion featuring natural language querying and vision-based document processing.",
     ["87% faster report generation (14.7s → 1.9s)", "Critical data pipeline errors resolved", "Natural language querying of fleet data", "Vision-based document processing", "ReAct agent with safety guardrails"],
     [".NET", "Entity Framework", "SQL Server", "Python", "Vision AI", "NLP"]),
    ("Financial Process Automation",
     "AI-Powered Financial Automation Platform",
     "Financial service providers managing multiple clients were drowning in repetitive work: invoice data entry, transaction categorization, and report generation. Each client had unique requirements, making standardized automation difficult.",
     "What began as a master's thesis exploring ML in accounting evolved into a comprehensive platform. The system combines AI-powered invoice extraction via custom OCR pipelines, automated transaction categorization, anomaly detection, and intelligent reporting. Built as multi-tenant SaaS with isolated data and configurable workflows.",
     ["AI-powered invoice extraction (OCR + NLP)", "ML-based transaction categorization", "Automated anomaly detection", "Multi-tenant architecture with data isolation", "Academic research foundation with industry validation"],
     ["Python", "FastAPI", "Next.js", "PostgreSQL", "scikit-learn", "Tesseract OCR", "Redis"]),
    ("SME Business Automation",
     "AI-Assisted Offer & Quote Creation for SMEs",
     "Small and medium enterprises spend disproportionate time creating offers and quotes. Each proposal requires pulling together descriptions, calculating pricing, applying client-specific terms — a process prone to inconsistency.",
     "We developed an AI-assisted offer creation tool for SMEs. The system guides users through a structured workflow, suggests descriptions and pricing based on historical patterns, auto-populates client data, and generates professional branded PDF proposals. AI assists from recommending line items to optimizing pricing.",
     ["Guided offer creation workflow", "AI-powered pricing suggestions", "Professional branded PDF generation", "Client history and pattern learning", "Designed for SMEs without sales teams"],
     ["Python", "FastAPI", "React", "NLP", "PDF Generation", "ML"]),
    ("Data Analytics & Monitoring",
     "Operational Intelligence & Predictive Monitoring",
     "Organizations managing distributed infrastructure lacked real-time visibility into performance. Data from sensors and systems sat in silos, making it impossible to detect emerging issues before they became costly failures.",
     "We built a monitoring platform that ingests operational data from multiple sources in real time, applies ML for anomaly detection and trend forecasting, and presents actionable insights through interactive dashboards. The platform supports configurable alerting, predictive maintenance, and automated reporting.",
     ["Real-time data ingestion and analysis", "ML-based anomaly detection and alerting", "Predictive maintenance forecasting", "Interactive dashboards with drill-down", "Scalable multi-site architecture"],
     ["Python", "Time-Series DB", "ML", "FastAPI", "React", "Docker", "IoT"]),
    ("AI Agent Orchestration",
     "AI Agent Deployment & Orchestration Platform",
     "Organizations adopting AI assistants face a scaling challenge: each agent requires manual setup, credential management, and infrastructure provisioning. Managing dozens of agents across teams becomes an operational burden.",
     "We built a mission control platform for deploying and orchestrating multiple AI agents as isolated Docker containers. One-click deployment with auto-configuration, secure credential vault with AES-256 encryption, intelligent reverse proxy with WebSocket bridging, and real-time monitoring across all agents.",
     ["One-click agent deployment with full isolation", "Secure credential vault (AES-256)", "Real-time mission control dashboard", "Intelligent reverse proxy with WebSocket", "Scalable to 25+ concurrent agents per host"],
     ["TypeScript", "Node.js", "Express", "React", "Docker", "SQLite", "WebSocket", "TLS"]),
]

def remove_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'  <w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

def pad_cell(cell, top=100, left=160, bottom=100, right=160):
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()
    margins = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'  <w:top w:w="{top}" w:type="dxa"/>'
        f'  <w:left w:w="{left}" w:type="dxa"/>'
        f'  <w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'  <w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(margins)

total = len(cases)
for idx, (tag, title, challenge, solution, results, techs) in enumerate(cases):
    doc.add_page_break()
    add_spacer(2)

    # Case study header
    add_text(f"Case Study {idx+1:02d} / {total:02d}", size=9, color=MUTED_COL, space_after=2)
    add_text(title, size=22, color=WHITE, bold=True, space_after=2)
    add_multirun([(tag, 11, ACCENT, False, True)], space_after=4)
    add_accent_line()
    add_spacer(6)

    # Two-column layout: Challenge+Solution | Outcomes+Tech
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Left column: Challenge + Solution
    left = table.rows[0].cells[0]
    left.width = Cm(14)
    shade_cell(left, DARK)

    p = left.paragraphs[0]
    run = p.add_run("CHALLENGE")
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(ACCENT)
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.letter_spacing = Pt(1.5)

    p2 = left.add_paragraph()
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(12)
    p2.paragraph_format.line_spacing = 1.25
    run2 = p2.add_run(challenge)
    run2.font.size = Pt(10)
    run2.font.color.rgb = rgb(DIM_COL)
    run2.font.name = 'Calibri'

    p3 = left.add_paragraph()
    run3 = p3.add_run("SOLUTION")
    run3.font.size = Pt(9)
    run3.font.color.rgb = rgb(ACCENT)
    run3.font.bold = True
    run3.font.name = 'Calibri'
    run3.font.letter_spacing = Pt(1.5)

    p4 = left.add_paragraph()
    p4.paragraph_format.space_before = Pt(4)
    p4.paragraph_format.line_spacing = 1.25
    run4 = p4.add_run(solution)
    run4.font.size = Pt(10)
    run4.font.color.rgb = rgb(DIM_COL)
    run4.font.name = 'Calibri'

    pad_cell(left, 60, 0, 60, 200)

    # Right column: Outcomes + Tech (in a panel)
    right = table.rows[0].cells[1]
    right.width = Cm(10.5)
    shade_cell(right, PANEL)

    p = right.paragraphs[0]
    run = p.add_run("OUTCOMES")
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(ACCENT)
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.letter_spacing = Pt(1.5)

    for r in results:
        pr = right.add_paragraph()
        pr.paragraph_format.space_before = Pt(3)
        pr.paragraph_format.line_spacing = 1.2
        run_bullet = pr.add_run("▸  ")
        run_bullet.font.size = Pt(9)
        run_bullet.font.color.rgb = rgb(ACCENT_DIM)
        run_bullet.font.name = 'Calibri'
        run_text = pr.add_run(r)
        run_text.font.size = Pt(10)
        run_text.font.color.rgb = rgb(TEXT_COL)
        run_text.font.name = 'Calibri'

    # Tech section
    pt = right.add_paragraph()
    pt.paragraph_format.space_before = Pt(14)
    run = pt.add_run("TECHNOLOGIES")
    run.font.size = Pt(9)
    run.font.color.rgb = rgb(ACCENT)
    run.font.bold = True
    run.font.name = 'Calibri'
    run.font.letter_spacing = Pt(1.5)

    pt2 = right.add_paragraph()
    pt2.paragraph_format.space_before = Pt(4)
    pt2.paragraph_format.line_spacing = 1.4
    for i, tech in enumerate(techs):
        run = pt2.add_run(tech)
        run.font.size = Pt(9)
        run.font.color.rgb = rgb(DIM_COL)
        run.font.name = 'Calibri'
        if i < len(techs) - 1:
            sep = pt2.add_run("  ·  ")
            sep.font.size = Pt(9)
            sep.font.color.rgb = rgb(MUTED_COL)
            sep.font.name = 'Calibri'

    pad_cell(right, 120, 180, 120, 180)
    remove_table_borders(table)

# ════════════════════════════════════════
# TECH STACK PAGE
# ════════════════════════════════════════
doc.add_page_break()
add_spacer(4)
add_section_title("Technology Stack", 26)
add_spacer(8)

stacks = [
    ("AI & Machine Learning", "Python  ·  TensorFlow  ·  PyTorch  ·  scikit-learn  ·  Hugging Face  ·  LangChain  ·  OpenAI  ·  Claude"),
    ("Natural Language Processing", "spaCy  ·  NLTK  ·  Transformers  ·  RAG Pipelines  ·  ChromaDB  ·  Pinecone  ·  Semantic Search"),
    ("Backend Development", "FastAPI  ·  Django  ·  ASP.NET Core  ·  Entity Framework  ·  Node.js  ·  Express"),
    ("Frontend & Visualization", "React  ·  Next.js  ·  Vue.js  ·  Nuxt  ·  Tailwind CSS  ·  D3.js  ·  Plotly"),
    ("Databases", "PostgreSQL  ·  SQL Server  ·  MariaDB  ·  Redis  ·  MongoDB  ·  Elasticsearch"),
    ("Infrastructure", "Docker  ·  Linux  ·  Nginx  ·  CI/CD  ·  Azure DevOps  ·  GitHub Actions"),
    ("Document Processing", "Tesseract OCR  ·  pdf2image  ·  Apache Tika  ·  Custom Pipelines"),
    ("Integration", "REST  ·  GraphQL  ·  WebSocket  ·  IMAP/SMTP  ·  OAuth 2.0  ·  SSO"),
]

table = doc.add_table(rows=len(stacks), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, (cat, items) in enumerate(stacks):
    cell0 = table.rows[i].cells[0]
    cell0.width = Cm(6)
    shade_cell(cell0, PANEL)
    p = cell0.paragraphs[0]
    run = p.add_run(cat)
    run.font.size = Pt(11)
    run.font.color.rgb = rgb(ACCENT)
    run.font.bold = True
    run.font.name = 'Calibri'
    pad_cell(cell0, 100, 160, 100, 80)

    cell1 = table.rows[i].cells[1]
    cell1.width = Cm(18.5)
    shade_cell(cell1, PANEL)
    p = cell1.paragraphs[0]
    run = p.add_run(items)
    run.font.size = Pt(10)
    run.font.color.rgb = rgb(DIM_COL)
    run.font.name = 'Calibri'
    pad_cell(cell1, 100, 80, 100, 160)

# Remove borders, add slight gap via cell margins
remove_table_borders(table)

# Small logo
add_spacer(20)
try:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(LOGO_W, width=Cm(3))
except:
    pass

# ════════════════════════════════════════
# CONTACT PAGE
# ════════════════════════════════════════
doc.add_page_break()

add_spacer(30)

# Logo
try:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run()
    run.add_picture(LOGO_W, width=Cm(4.5))
except:
    pass

add_spacer(10)

add_text("Let's Build Something Intelligent Together", size=24, color=ACCENT, bold=True, align='center', space_after=4)
add_accent_line()
add_spacer(6)
add_text("Get in touch for a free initial consultation.", size=13, color=TEXT_COL, align='center', space_after=24)

# Founders signature block — two columns
add_spacer(4)
table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Gergo
left = table.rows[0].cells[0]
shade_cell(left, PANEL)
p = left.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Gergo Kiss, MSc")
run.font.size = Pt(14)
run.font.color.rgb = rgb(WHITE)
run.font.bold = True
run.font.name = 'Calibri'
p2 = left.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Managing Director  ·  AI Systems Architect")
run2.font.size = Pt(10)
run2.font.color.rgb = rgb(ACCENT)
run2.font.name = 'Calibri'
p3 = left.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(4)
run3 = p3.add_run("gergo.kiss@y2k.global")
run3.font.size = Pt(9.5)
run3.font.color.rgb = rgb(DIM_COL)
run3.font.name = 'Calibri'
pad_cell(left, 140, 120, 140, 120)

# Rami
right = table.rows[0].cells[1]
shade_cell(right, PANEL)
p = right.paragraphs[0]
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Ing. Ramazan Yildirim")
run.font.size = Pt(14)
run.font.color.rgb = rgb(WHITE)
run.font.bold = True
run.font.name = 'Calibri'
p2 = right.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = p2.add_run("Managing Director  ·  AI Solutions Engineer")
run2.font.size = Pt(10)
run2.font.color.rgb = rgb(ACCENT)
run2.font.name = 'Calibri'
p3 = right.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
p3.paragraph_format.space_before = Pt(4)
run3 = p3.add_run("ramazan.yildirim@y2k.global")
run3.font.size = Pt(9.5)
run3.font.color.rgb = rgb(DIM_COL)
run3.font.name = 'Calibri'
pad_cell(right, 140, 120, 140, 120)

remove_table_borders(table)

add_spacer(8)
add_text("+43 1 442 20 143  ·  office@y2k.global  ·  www.y2k.global", size=11, color=TEXT_COL, align='center', space_after=14)

add_muted_line()
add_spacer(6)

# Company details - two column
table = doc.add_table(rows=3, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

info = [
    [("Y2K Global OÜ", 11, WHITE, True), ("In cooperation with", 9, MUTED_COL, False)],
    [("Sepapaja 6  ·  15551 Tallinn  ·  Estonia", 9.5, MUTED_COL, False), ("KISS IT Solutions e.U.", 11, WHITE, True)],
    [("Reg: 17373395  ·  VAT: EE102928187", 9.5, MUTED_COL, False), ("Seitenstettengasse 5/37  ·  1010 Vienna  ·  Austria\nReg: FN606200x  ·  VAT: ATU68895724", 9.5, MUTED_COL, False)],
]

for i, row_data in enumerate(info):
    for j, (text, size, color, bold) in enumerate(row_data):
        cell = table.rows[i].cells[j]
        shade_cell(cell, DARK)
        p = cell.paragraphs[0]
        if j == 0:
            pass
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for line_idx, line in enumerate(text.split('\n')):
            if line_idx > 0:
                p = cell.add_paragraph()
                if j == 1:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(line)
            run.font.size = Pt(size)
            run.font.color.rgb = rgb(color)
            run.font.bold = bold
            run.font.name = 'Calibri'
        pad_cell(cell, 40, 80, 40, 80)

remove_table_borders(table)

# ── Save ──
doc.save(OUT)
sz = os.path.getsize(OUT)
print(f"Done: {sz:,} bytes ({sz/1024:.1f} KB)")
